import io
import os
import json
import logging
import re
import docx
import pypdf
import google.generativeai as genai

from reportlab.lib.pagesizes import letter, A4
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable, KeepTogether
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

logger = logging.getLogger(__name__)

# Register Arial / Times New Roman font for Vietnamese support if on Windows
FONT_NAME = "Helvetica"
FONT_NAME_BOLD = "Helvetica-Bold"

try:
    windows_fonts = "C:/Windows/Fonts"
    arial_path = os.path.join(windows_fonts, "arial.ttf")
    arial_bd_path = os.path.join(windows_fonts, "arialbd.ttf")
    if os.path.exists(arial_path) and os.path.exists(arial_bd_path):
        pdfmetrics.registerFont(TTFont("Arial", arial_path))
        pdfmetrics.registerFont(TTFont("Arial-Bold", arial_bd_path))
        FONT_NAME = "Arial"
        FONT_NAME_BOLD = "Arial-Bold"
        logger.info("Successfully registered Arial font for PDF generation.")
except Exception as e:
    logger.warning(f"Could not register Arial font: {e}")

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if GEMINI_API_KEY:
    genai.configure(api_key=GEMINI_API_KEY)

def _clean_uuid_text(text: str) -> str:
    """Thay thế triệt để mã UUID thô bằng tên gọi Thành viên ngắn gọn."""
    if not isinstance(text, str):
        return text
    uuid_pattern = r'[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{4}-[0-9a-fA-F]{12}'
    return re.sub(uuid_pattern, "Thành viên", text)

def _sanitize_data_uuids(obj):
    if isinstance(obj, str):
        return _clean_uuid_text(obj)
    elif isinstance(obj, list):
        return [_sanitize_data_uuids(item) for item in obj]
    elif isinstance(obj, dict):
        return {k: _sanitize_data_uuids(v) for k, v in obj.items()}
    return obj

def parse_custom_template(file_bytes: bytes, filename: str) -> str:
    """
    Trích xuất văn bản dàn ý từ tệp mẫu .docx hoặc .pdf do người dùng tải lên.
    """
    ext = os.path.splitext(filename)[1].lower()
    text_content = ""
    
    if ext == ".docx":
        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    paragraphs.append(row_text)
        text_content = "\n".join(paragraphs)
    elif ext == ".pdf":
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pages_text = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pages_text.append(t)
        text_content = "\n".join(pages_text)
    else:
        raise ValueError("Định dạng tệp không được hỗ trợ. Vui lòng tải lên tệp .docx hoặc .pdf")

    if not text_content.strip():
        raise ValueError("Không thể đọc được nội dung từ tệp mẫu đã chọn.")
        
    return text_content.strip()

GEMINI_FALLBACK_MODELS = [
    'gemini-3.5-flash',
    'gemini-3.5-flash-lite',
    'gemini-3.1-flash-lite',
    'gemini-3-flash-preview',
    'gemini-2.5-pro',
    'gemini-2.0-flash'
]

def _call_gemini_with_model_fallback(prompt: str, max_retries_per_model=2, delay=2):
    import time
    last_exception = None
    for model_name in GEMINI_FALLBACK_MODELS:
        logger.info(f"Attempting report generation using Gemini model: '{model_name}'...")
        try:
            model = genai.GenerativeModel(model_name)
            for attempt in range(max_retries_per_model):
                try:
                    response = model.generate_content(prompt)
                    if response and response.text:
                        logger.info(f"Successfully generated report using model: '{model_name}'.")
                        return response.text
                except Exception as e:
                    err_str = str(e)
                    if "429" in err_str or "quota" in err_str.lower() or "ResourceExhausted" in err_str:
                        logger.warning(f"Model '{model_name}' rate limited (429). Retrying in {delay}s...")
                        time.sleep(delay)
                    else:
                        raise e
        except Exception as e:
            last_exception = e
            logger.warning(f"Model '{model_name}' failed with error: {e}. Falling back to next available model in chain...")
            continue
            
    raise RuntimeError(f"All Gemini fallback models exhausted for report generation: {last_exception}")

def generate_report_content(meeting_title: str, segments: list, speaker_stats: list, template_type: str = "default", custom_template_text: str = None) -> dict:
    if not GEMINI_API_KEY:
        raise RuntimeError("GEMINI_API_KEY chưa được cấu hình.")

    formatted_transcript = ""
    for seg in segments:
        spk = seg.get("speaker", seg.get("speaker_id", "Thành viên"))
        text = seg.get("text", seg.get("content", ""))
        start = seg.get("start", seg.get("start_time", 0))
        m_s = f"{int(start//60):02d}:{int(start%60):02d}"
        formatted_transcript += f"[{m_s}] {spk}: {text}\n"

    # Speaker summary list
    speakers_summary_str = ", ".join([spk.get("name", "Thành viên") for spk in speaker_stats if spk.get("name")])

    if template_type == "custom" and custom_template_text:
        prompt = f"""
Bạn là một Trợ lý AI phân tích và biên soạn Biên Bản Cuộc Họp cao cấp.
Nhiệm vụ: Lập báo cáo cuộc họp dựa trên MẪU YÊU CẦU DÀN Ý TỦY CHỈNH dưới đây.

RÀNG BUỘC TUYỆT ĐỐI BẮT BUỘC:
1. KHÔNG ĐƯỢC để mã ID/UUID thô (như c0f65525-750c-465c-820f-de62793f9074). Luôn gọi tên thành viên bằng TÊN ĐÀNG HOÀNG từ bản ghi (như Đức cớp, Minh vẽ, Khánh, hoặc Thành viên A/B).
2. KHÔNG ĐƯỢC thêm, bịa đặt hoặc tự suy đoán bất kỳ thông tin nào ngoài bản ghi âm cuộc họp (100% tri thức khai thác từ transcript).
3. Phân tích kĩ mẫu yêu cầu và điền thông tin phù hợp nhất theo transcript.

MẪU YÊU CẦU DÀN Ý CỦA NGƯỜI DÙNG:
{custom_template_text}

BẢN GHI ÂM CUỘC HỌP (TRANSCRIPT):
{formatted_transcript}

Yêu cầu Output: Trả về CHỈ duy nhất 1 JSON Object hợp lệ (không chứa markdown ```json):
{{
  "title": "{meeting_title}",
  "sections": [
    {{
      "heading": "Tên mục theo mẫu dàn ý",
      "content": "Nội dung chi tiết tổng hợp từ cuộc họp cho mục này"
    }}
  ]
}}
"""
    else:
        prompt = f"""
Bạn là một Trợ lý AI biên soạn BIÊN BẢN HỌP CHUẨN DOANH NGHIỆP VIỆT NAM.
Nhiệm vụ: Tổng hợp Biên Bản Cuộc Họp theo Tiêu chuẩn Biên bản họp chuyên nghiệp tại Việt Nam.

RÀNG BUỘC CHUẨN MỰC BẮT BUỘC:
1. TUYỆT ĐỐI KHÔNG ĐƯỢC tự ý đặt tên thành "Thành viên A", "Thành viên B" hay "Thành viên 1/2". PHẦI SỬ DỤNG CHÍNH XÁC TÊN THÀNH VIÊN THỰC TẾ TRONG DANH SÁCH BẢN GHI ({speakers_summary_str}). Trích xuất và ghi chính xác tên người phát biểu từng ý kiến.
2. KHÔNG ĐƯỢC để bất kỳ mã ID/UUID thô dạng chuỗi ký tự dài (như c0f65525-750c-465c-820f-de62793f9074).
3. KHÔNG ĐƯỢC thêm, bịa đặt hoặc tự suy đoán bất kỳ thông tin nào ngoài bản ghi âm cuộc họp (100% tri thức khai thác từ transcript).
4. Nội dung rõ ràng, mạch lạc, hành văn chuẩn mực tiếng Việt công sở.

BẢN GHI ÂM CUỘC HỌP (TRANSCRIPT):
{formatted_transcript}

DANH SÁCH THÀNH VIÊN THAM GIA:
{speakers_summary_str}

Yêu cầu Output: Trả về CHỈ duy nhất 1 JSON Object hợp lệ (không chứa markdown ```json):
{{
  "title": "{meeting_title}",
  "overview": "Mục tiêu và tóm tắt tổng quan cuộc họp trong 2-3 câu ngắn gọn.",
  "attendees_summary": "Cuộc họp diễn ra với sự tham gia thảo luận của các thành viên: {speakers_summary_str}.",
  "key_discussions": [
    {{
      "topic": "Tên chủ đề thảo luận 1",
      "detail": "Nội dung phân tích thảo luận chi tiết giữa các thành viên, ghi rõ tên thành viên phát biểu."
    }}
  ],
  "decisions": [
    "Kết luận hoặc quyết định đã thống nhất 1",
    "Kết luận hoặc quyết định đã thống nhất 2"
  ],
  "action_items": [
    {{
      "task": "Nhiệm vụ cần thực hiện",
      "assignee": "Tên thành viên phụ trách (sử dụng tên thật từ danh sách {speakers_summary_str}, hoặc Cả nhóm)",
      "deadline": "Thời hạn / Ghi chú"
    }}
  ],
  "assistant_notes": "Ghi chú khuyến nghị thêm của Trợ lý AI dựa trên diễn biến cuộc họp."
}}
"""

    text_resp = _call_gemini_with_model_fallback(prompt).strip()
    if text_resp.startswith("```json"):
        text_resp = text_resp[7:]
    elif text_resp.startswith("```"):
        text_resp = text_resp[3:]
    if text_resp.endswith("```"):
        text_resp = text_resp[:-3]

    try:
        data = json.loads(text_resp.strip())
        data = _sanitize_data_uuids(data)
        return data
    except Exception as e:
        logger.error(f"Lỗi parse JSON báo cáo từ Gemini: {e}. Output gốc: {text_resp}")
        fallback_data = {
            "title": meeting_title,
            "overview": "Báo cáo cuộc họp được tổng hợp tự động.",
            "attendees_summary": f"Cuộc họp có sự tham gia của: {speakers_summary_str}.",
            "key_discussions": [{"topic": "Nội dung cuộc họp", "detail": text_resp}],
            "decisions": ["Hoàn tất trao đổi trong cuộc họp."],
            "action_items": [],
            "assistant_notes": "Cuộc họp đã được ghi nhận."
        }
        return _sanitize_data_uuids(fallback_data)

class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_number(num_pages)
            canvas.Canvas.showPage(self)
        canvas.Canvas.save(self)

    def draw_page_number(self, page_count):
        self.saveState()
        self.setFont(FONT_NAME, 8)
        self.setFillColor(colors.HexColor("#64748b"))
        
        # Header text on subsequent pages
        if self._pageNumber > 1:
            self.drawString(36, 810, "FLOW MEETING INTELLIGENCE — HỆ THỐNG QUẢN TRỊ CUỘC HỌP")
            self.setStrokeColor(colors.HexColor("#e2e8f0"))
            self.setLineWidth(0.5)
            self.line(36, 802, 559, 802)

        # Footer line
        self.setStrokeColor(colors.HexColor("#cbd5e1"))
        self.setLineWidth(0.5)
        self.line(36, 45, 559, 45)

        # Footer text
        self.drawString(36, 30, "Biên bản cuộc họp được khởi tạo tự động bởi Hệ thống Flow AI")
        page_text = f"Trang {self._pageNumber} / {page_count}"
        self.drawRightString(559, 30, page_text)
        self.restoreState()

def build_pdf_report(report_data: dict, meeting_date=None) -> bytes:
    """
    Xuất báo cáo cuộc họp thành file PDF Biên Bản Họp chuẩn Việt Nam.
    """
    report_data = _sanitize_data_uuids(report_data)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=36,
        rightMargin=36,
        topMargin=45,
        bottomMargin=55
    )

    styles = getSampleStyleSheet()

    # Define Vietnamese Meeting Minutes styles
    agency_style = ParagraphStyle(
        'AgencyHeader',
        parent=styles['Normal'],
        fontName=FONT_NAME_BOLD,
        fontSize=9.5,
        leading=13,
        alignment=0, # Left
        textColor=colors.HexColor("#0f172a")
    )

    national_header_title = ParagraphStyle(
        'NationalHeaderTitle',
        parent=styles['Normal'],
        fontName=FONT_NAME,
        fontSize=9.5,
        leading=13,
        alignment=1, # Centered
        textColor=colors.HexColor("#0f172a")
    )

    doc_title_style = ParagraphStyle(
        'DocTitleCustom',
        parent=styles['Title'],
        fontName=FONT_NAME_BOLD,
        fontSize=20,
        leading=24,
        alignment=1, # Centered
        textColor=colors.HexColor("#4c1d95"),
        spaceAfter=4
    )

    doc_subtitle_style = ParagraphStyle(
        'DocSubtitleCustom',
        parent=styles['Normal'],
        fontName=FONT_NAME,
        fontSize=11,
        leading=15,
        alignment=1, # Centered
        textColor=colors.HexColor("#334155"),
        spaceAfter=14
    )

    h1_style = ParagraphStyle(
        'H1Custom',
        parent=styles['Heading1'],
        fontName=FONT_NAME_BOLD,
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#5b21b6"),
        spaceBefore=12,
        spaceAfter=6,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'BodyCustom',
        parent=styles['Normal'],
        fontName=FONT_NAME,
        fontSize=9.5,
        leading=14.5,
        alignment=4,
        textColor=colors.HexColor("#1e293b"),
        spaceAfter=5
    )

    bold_body_style = ParagraphStyle(
        'BoldBodyCustom',
        parent=body_style,
        fontName=FONT_NAME_BOLD,
        textColor=colors.HexColor("#0f172a")
    )

    signature_title_style = ParagraphStyle(
        'SigTitle',
        parent=styles['Normal'],
        fontName=FONT_NAME_BOLD,
        fontSize=10.5,
        leading=14,
        alignment=1,
        textColor=colors.HexColor("#0f172a")
    )

    signature_sub_style = ParagraphStyle(
        'SigSub',
        parent=styles['Normal'],
        fontName=FONT_NAME,
        fontSize=9,
        leading=12,
        alignment=1,
        textColor=colors.HexColor("#64748b")
    )

    story = []

    # Format date dynamically
    now = meeting_date if meeting_date else datetime.now()
    date_str = f"Ngày {now.day:02d} tháng {now.month:02d} năm {now.year}"

    # 1. Header Block (Quốc hiệu & Tiêu ngữ chuẩn Việt Nam)
    left_header = Paragraph("<b>HỆ THỐNG FLOW INTELLIGENCE</b><br/>CƠ QUAN / ĐƠN VỊ CHỦ QUẢN", agency_style)
    right_header = Paragraph(
        "<b>CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM</b><br/>"
        "<b><u>Độc lập - Tự do - Hạnh phúc</u></b><br/>"
        f"<font size=8.5 color='#475569'><i>{date_str}</i></font>",
        national_header_title
    )
    
    header_table = Table([[left_header, right_header]], colWidths=[240, 283])
    header_table.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('ALIGN', (0,0), (0,0), 'LEFT'),
        ('ALIGN', (1,0), (1,0), 'CENTER'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 8),
    ]))
    story.append(header_table)
    story.append(Spacer(1, 8))

    # 2. Document Main Title
    title_text = report_data.get("title", "CUỘC HỌP THẢO LUẬN")
    story.append(Paragraph("BIÊN BẢN HỌP", doc_title_style))
    story.append(Paragraph(f"Về việc: {title_text}", doc_subtitle_style))
    story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#7c3aed"), spaceAfter=14))

    # 3. Content Sections
    if "sections" in report_data and isinstance(report_data["sections"], list):
        for sec in report_data["sections"]:
            heading = sec.get("heading", "Mục báo cáo")
            content = sec.get("content", "")
            story.append(Paragraph(heading.upper(), h1_style))
            story.append(Paragraph(content.replace("\n", "<br/>"), body_style))
            story.append(Spacer(1, 8))
    else:
        # Official Vietnamese Meeting Minutes Flow (Biên bản họp chuẩn Việt Nam)
        # Section I: Thành phần tham dự
        story.append(Paragraph("I. THỜI GIAN & THÀNH PHẦN THAM DỰ", h1_style))
        if report_data.get("attendees_summary"):
            story.append(Paragraph(report_data["attendees_summary"], body_style))
        else:
            story.append(Paragraph("Cuộc họp diễn ra với sự tham gia thảo luận của các thành viên trong dự án.", body_style))
        story.append(Spacer(1, 6))

        # Section II: Nội dung thảo luận chi tiết
        story.append(Paragraph("II. NỘI DUNG THẢO LUẬN CHI TIẾT", h1_style))
        if report_data.get("overview"):
            story.append(Paragraph(f"<b>1. Tổng quan bối cảnh:</b> {report_data['overview']}", body_style))
            story.append(Spacer(1, 4))

        discussions = report_data.get("key_discussions", [])
        if discussions:
            story.append(Paragraph("<b>2. Chi tiết các nội dung trao đổi:</b>", body_style))
            for i, disc in enumerate(discussions, 1):
                topic = disc.get("topic", f"Chủ đề {i}")
                detail = disc.get("detail", "")
                story.append(Paragraph(f"<b>2.{i}. {topic}</b>", bold_body_style))
                story.append(Paragraph(detail.replace("\n", "<br/>"), body_style))
                story.append(Spacer(1, 4))

        # Section III: Quyết định & Kết luận
        decisions = report_data.get("decisions", [])
        if decisions:
            story.append(Paragraph("III. QUYẾT ĐỊNH & KẾT LUẬN THỐNG NHẤT", h1_style))
            for dec in decisions:
                story.append(Paragraph(f"• {dec}", body_style))
            story.append(Spacer(1, 6))

        # Section IV: Phân công công việc (Action Items Table)
        action_items = report_data.get("action_items", [])
        if action_items:
            story.append(Paragraph("IV. KẾ HOẠCH & PHÂN CÔNG CÔNG VIỆC (ACTION ITEMS)", h1_style))
            
            table_data = [[
                Paragraph("<b>STT</b>", bold_body_style),
                Paragraph("<b>Nhiệm vụ cụ thể</b>", bold_body_style),
                Paragraph("<b>Người phụ trách</b>", bold_body_style),
                Paragraph("<b>Thời hạn / Ghi chú</b>", bold_body_style)
            ]]
            
            for idx, item in enumerate(action_items, 1):
                task = item.get("task", "")
                assignee = item.get("assignee", "Chưa phân công")
                deadline = item.get("deadline", "---")
                table_data.append([
                    Paragraph(str(idx), body_style),
                    Paragraph(task, body_style),
                    Paragraph(assignee, body_style),
                    Paragraph(deadline, body_style)
                ])
                
            task_table = Table(table_data, colWidths=[30, 230, 130, 133])
            task_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#f1f5f9")),
                ('TEXTCOLOR', (0,0), (-1,0), colors.HexColor("#0f172a")),
                ('ALIGN', (0,0), (0,-1), 'CENTER'),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                ('TOPPADDING', (0,0), (-1,-1), 6),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#cbd5e1")),
            ]))
            story.append(task_table)
            story.append(Spacer(1, 10))

        if report_data.get("assistant_notes"):
            story.append(Paragraph("V. GHI CHÚ KHUYẾN NGHỊ TỪ TRỢ LÝ AI", h1_style))
            story.append(Paragraph(report_data["assistant_notes"], body_style))
            story.append(Spacer(1, 10))

    # 4. Signature Block (Phần ký tên chuẩn Biên Bản Họp Việt Nam)
    story.append(KeepTogether([
        Spacer(1, 15),
        Paragraph("<i>Biên bản họp được các thành viên nhất trí thông qua và có hiệu lực kể từ ngày ký./.</i>", ParagraphStyle('SubText', parent=body_style, fontSize=9, fontName=FONT_NAME, textColor=colors.HexColor("#475569"), spaceAfter=15)),
        Table([
            [
                Paragraph("<b>THƯ KÝ</b>", signature_title_style),
                Paragraph("<b>CHỦ TỌA / CÁC THÀNH VIÊN</b>", signature_title_style)
            ],
            [
                Paragraph("<i>(Ký, ghi rõ họ tên)</i>", signature_sub_style),
                Paragraph("<i>(Ký, ghi rõ họ tên)</i>", signature_sub_style)
            ],
            [
                Spacer(1, 45),
                Spacer(1, 45)
            ]
        ], colWidths=[260, 263], style=[
            ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ])
    ]))

    doc.build(story, canvasmaker=NumberedCanvas)
    buffer.seek(0)
    return buffer.getvalue()
